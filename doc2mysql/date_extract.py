import docx
import datetime
import os
import re
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))




def b27_date_extract(FiledocxPathName,project_name,filename):
    doc = docx.Document(FiledocxPathName)
    patient_name = doc.tables[0].rows[0].cells[1].text
    # 去除姓名中空格
    patient_name = patient_name.replace(' ', '')
    patient_sex = doc.tables[0].rows[0].cells[3].text
    patient_age = doc.tables[0].rows[0].cells[5].text
    patient_department = doc.tables[0].rows[0].cells[7].text
    patient_ID = doc.tables[0].rows[1].cells[3].text
    submit_doctor = doc.tables[0].rows[1].cells[7].text
    # # print("%s 信息栏提取结束" % FiledocxPathName)
    # 提取result及opinion
    ss = []
    report_content = doc.paragraphs
    for i in report_content:
        ss.append(i.text)
    patient_result = '无法获取检查结果，请联系相关科室'
    for l in ss:
        if '阴性' in l:
            patient_result = '阴性'
        elif '阳性' in l:
            patient_result = '阳性'
        else:
            pass
    # # print(patient_result)
    report_opinion = ''
    for m in ss:
        if '建议复查' in m:
            report_opinion = '注：接近临界值，建议复查\t'
        else:
            pass
    # # print(report_opinion)
    # 报告生成时间
    DocFilePathName = os.path.join(BASE_DIR, 'upload', project_name, filename)
    time1 = os.path.getmtime(DocFilePathName)
    time2 = datetime.datetime.fromtimestamp(time1)
    RPmtime = time2
    RPmtime_urltime = datetime.datetime.strftime(RPmtime, "%Y-%m-%d-%H-%M-%S-%f")
    # # print(RPmtime)
    # # print("%s 报告生成时间提取结束" % FiledocxPathName)
    # 报告时间及检测报告人员
    try:
        Submit_Sample_time = doc.tables[1].rows[0].cells[1].text
    except:
        Submit_Sample_time = "未知"
    try:
        Inspection_person = doc.tables[1].rows[0].cells[3].text
    except:
        Inspection_person = "未知"
    try:
        Detection_time = doc.tables[1].rows[1].cells[1].text
    except:
        Detection_time = "未知"
    try:
        Report_Doctor = doc.tables[1].rows[1].cells[3].text
    except:
        Report_Doctor = "未知"
    # # print(Submit_Sample_time, Inspection_person, Detection_time, Report_Doctor)
    # # print("~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~")
    # # print("%s 报告人提取结束" % FiledocxPathName)
    # 报告生成时间
    DocFilePathName = os.path.join(BASE_DIR, 'upload', project_name, filename)
    time1 = os.path.getmtime(DocFilePathName)
    time2 = datetime.datetime.fromtimestamp(time1)
    RPmtime = time2
    # # print(RPmtime)
    # # print("%s 报告生成时间提取结束" % FiledocxPathName)
    b27_date = {
        'patient_name': patient_name,
        'patient_sex': patient_sex,
        'patient_age': patient_age,
        'patient_department': patient_department,
        'patient_ID': patient_ID,
        'submit_doctor': submit_doctor,
        'patient_result': patient_result,
        'report_opinion': report_opinion,
        'RPmtime': RPmtime,
        'Submit_Sample_time': Submit_Sample_time,
        'Detection_time': Detection_time,
        'Inspection_person': Inspection_person,
        'Report_Doctor': Report_Doctor,
    }
    return b27_date





def TH_date_extract(FiledocxPathName,project_name,filename):
    doc = docx.Document(FiledocxPathName)
    patient_name = doc.tables[0].rows[0].cells[1].text
    # 去除姓名中空格
    patient_name = patient_name.replace(' ', '')
    patient_sex = doc.tables[0].rows[0].cells[3].text
    patient_age = doc.tables[0].rows[0].cells[5].text
    patient_department = doc.tables[0].rows[0].cells[7].text
    patient_ID = doc.tables[0].rows[1].cells[3].text
    submit_doctor = doc.tables[0].rows[1].cells[7].text
    # print("%s 信息栏提取结束" % FiledocxPathName)
    # 提取th 数据
    patient_IL2 = doc.tables[1].rows[1].cells[1].text
    patient_IL4 = doc.tables[1].rows[2].cells[1].text
    patient_IL6 = doc.tables[1].rows[3].cells[1].text
    patient_IL10 = doc.tables[1].rows[4].cells[1].text
    patient_TNF = doc.tables[1].rows[5].cells[1].text
    patient_IFN = doc.tables[1].rows[6].cells[1].text
    # 报告生成时间
    DocFilePathName = os.path.join(BASE_DIR, 'upload', project_name, filename)
    time1 = os.path.getmtime(DocFilePathName)
    time2 = datetime.datetime.fromtimestamp(time1)
    RPmtime = time2
    # print("%s 报告生成时间提取结束" % FiledocxPathName)
    # 报告时间及检测报告人员
    try:
        Submit_Sample_time = doc.tables[2].rows[0].cells[1].text
    except:
        Submit_Sample_time = "未知"
    try:
        Inspection_person = doc.tables[2].rows[0].cells[3].text
    except:
        Inspection_person = "未知"
    try:
        Detection_time = doc.tables[2].rows[1].cells[1].text
    except:
        Detection_time = "未知"
    try:
        Report_Doctor = doc.tables[2].rows[1].cells[3].text
    except:
        Report_Doctor = "未知"
    # print(Submit_Sample_time, Inspection_person, Detection_time, Report_Doctor)
    # print("~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~~")
    # print("%s 报告人提取结束" % FiledocxPathName)
    # 报告生成时间
    DocFilePathName = os.path.join(BASE_DIR, 'upload', project_name, filename)
    time1 = os.path.getmtime(DocFilePathName)
    time2 = datetime.datetime.fromtimestamp(time1)
    RPmtime = time2
    # print(RPmtime)
    # print("%s 报告生成时间提取结束" % FiledocxPathName)
    TH_date = {
        'patient_name': patient_name,
        'patient_sex': patient_sex,
        'patient_age': patient_age,
        'patient_department': patient_department,
        'patient_ID': patient_ID,
        'submit_doctor': submit_doctor,
        'patient_IL2': patient_IL2,
        'patient_IL4': patient_IL4,
        'patient_IL6': patient_IL6,
        'patient_IL10': patient_IL10,
        'patient_TNF': patient_TNF,
        'patient_IFN': patient_IFN,
        'RPmtime': RPmtime,
        'Submit_Sample_time': Submit_Sample_time,
        'Detection_time': Detection_time,
        'Inspection_person': Inspection_person,
        'Report_Doctor': Report_Doctor,
    }
    return TH_date


